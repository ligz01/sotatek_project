import os
from paddleocr import PaddleOCR
from docx import Document

ocr = PaddleOCR(use_textline_orientation=True, lang='en')

doc = Document()
doc.add_heading('Extracted Notes', 0)

image_folder = 'outputs'  # 🔥 đúng path

files = os.listdir(image_folder)
print("FILES:", files)

for img_name in files:
    img_path = os.path.join(image_folder, img_name)

    if not img_name.endswith(('.png', '.jpg', '.jpeg')):
        continue

    print("Processing:", img_path)

    result = ocr.predict(img_path)

    for res in result:
        for text in res['rec_texts']:
            doc.add_paragraph(text)

doc.save('output.docx')
print("✅ DONE → output.docx")